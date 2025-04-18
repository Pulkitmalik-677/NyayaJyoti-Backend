from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pandas as pd
import torch
import re
from docx import Document
import tempfile
from transformers import AutoTokenizer, AutoModelForCausalLM
from sentence_transformers import SentenceTransformer, util
from clause_generator import find_best_match_semantic, build_prompt, generate_clause, fill_parameters_dynamic

app = Flask(__name__)
CORS(app)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
tokenizer = AutoTokenizer.from_pretrained("EleutherAI/gpt-neo-1.3B")
model = AutoModelForCausalLM.from_pretrained("EleutherAI/gpt-neo-1.3B")
if tokenizer.pad_token is None:
    tokenizer.add_special_tokens({'pad_token': '[PAD]'})
    model.resize_token_embeddings(len(tokenizer))
model = model.to(device)

embedder = SentenceTransformer("all-MiniLM-L6-v2")

prompt_df = pd.read_csv("datasets/Verified_50_Loan_Clause_Examples_final_fixed.csv")
prompt_df["parameters"] = prompt_df["parameters"].fillna("").astype(str)
prompt_df["parameters"] = prompt_df["parameters"].apply(
    lambda x: ", ".join(sorted(set(p.strip() for p in x.split(",") if p.strip().lower() != "nan")))
)
instruction_embeddings = embedder.encode(prompt_df["instruction"].tolist(), convert_to_tensor=True)

@app.route("/generate", methods=["POST"])
def generate():
    try:
        data = request.json
        doc = Document("templates/loan_agreement.docx")

        # Fill standard fields
        for key, value in data.items():
            if key != "Special_Clauses":
                for para in doc.paragraphs:
                    if f"[{key}]" in para.text:
                        original = para.text
                        para.text = para.text.replace(f"[{key}]", value)
                        print(f"📄 Updated Line:\n\nBefore: {original}\n\nAfter: {para.text}\n")

        # Special Clause Logic
        special_input = data.get("Special_Clauses", "").strip()
        if special_input:
            match_row = find_best_match_semantic(special_input, embedder, instruction_embeddings, prompt_df)
            if match_row is not None:
                prompt = build_prompt(match_row['example_1'], match_row['example_2'], match_row['instruction'])
                raw_clause = generate_clause(prompt, tokenizer, model, device)
                filled_clause, _ = fill_parameters_dynamic(raw_clause, match_row.get('parameters', ''))
                for para in doc.paragraphs:
                    if "[Special_Clauses]" in para.text:
                        para.text = para.text.replace("[Special_Clauses]", filled_clause)
                        print(f"📄 Inserted Special Clause: {para.text}")
                        break
        else:
            for para in doc.paragraphs:
                if "[Special_Clauses]" in para.text:
                    para.text = para.text.replace("[Special_Clauses]", "N/A")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return send_file(tmp.name, as_attachment=True, download_name="Loan_Agreement.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/", methods=["GET"])
def home():
    return "Nyaya Jyoti GPT-Neo Legal Clause Generator is running."

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
